import json
import random
from docx import Document

def loadJson(fname):
    with open(fname, 'r') as file:
        mcqs = json.load(file)
    return mcqs

def generateDoc(mcqs, fname):
    doc = Document()
    doc.add_heading('Multiple Choice Questions', level=1)
    
    for i, mcq in enumerate(mcqs):
        doc.add_paragraph(f"Question {i + 1}: {mcq['question']}")
        options = mcq['options']
        optionStr = " ".join([f"({chr(65 + j)}) {option} \t" for j, option in enumerate(options)])
        doc.add_paragraph(optionStr)
        doc.add_paragraph()  

    doc.save(fname)

def generate_answers_file(mcqs, fname):
    with open(fname, "w") as file:
        file.write("Answers\n\n")
        for i, mcq in enumerate(mcqs):
            file.write(f"Question {i + 1}: {mcq['correct_option']}\n")

def main():
    mcqs = loadJson('mcqs.json')
    
    if mcqs:
        allMcqs = []
        for i in range(4):
            newMcqs = mcqs[:]
            random.shuffle(newMcqs)
            docFname = f"questions{i}.docx"
            generateDoc(newMcqs, docFname)
            allMcqs.append((docFname, newMcqs))
        
        with open("answers.txt", "w") as ansW:
            for docFname, mcqs in allMcqs:
                ansW.write(f"Answers for {docFname}\n")
                for i, mcq in enumerate(mcqs):
                    ansW.write(f"Question {i + 1}: {mcq['correct_option']}\n")
                ansW.write("\n")
    else:
        print("dint receive mcqs")

if __name__ == "__main__":
    main()
